"""
将 Markdown（学术报告常用结构）导出为 Word .docx（不依赖 Pandoc）。
支持：多级标题、粗体 **、引用 >、无序列表 -、Markdown 表格、独立图片 ![](path)。
代码块与 Mermaid：以 Courier 小号等宽段落保留。
用法:
  python export_md_to_docx.py report_academic_integrated.md
  python export_md_to_docx.py report_academic_integrated.md -o out.docx --base-dir "e:\\本基中间结果"
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def strip_md_bold(cell_text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)


def add_paragraph_with_inline(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        m = re.match(r"^\*\*(.+)\*\*$", part)
        if m:
            p.add_run(m.group(1)).bold = True
        else:
            p.add_run(part)
    return p


def flush_table(doc: Document, rows: List[List[str]]) -> None:
    if len(rows) < 2:
        return
    header = rows[0]
    body_start = 1
    if body_start < len(rows) and all(re.match(r"^[\s\-:]*$", c) for c in rows[body_start]):
        body_start += 1
    data_rows = rows[body_start:] if body_start <= len(rows) else []
    ncols = max(len(header), max((len(r) for r in data_rows), default=1))
    while len(header) < ncols:
        header.append("")
    tbl = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    tbl.style = "Table Grid"
    for j, cell in enumerate(header):
        tbl.rows[0].cells[j].text = strip_md_bold(cell.strip())
    for i, row in enumerate(data_rows, start=1):
        for j in range(ncols):
            val = strip_md_bold(row[j].strip()) if j < len(row) else ""
            tbl.rows[i].cells[j].text = val
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path, base_dir: Path) -> None:
    text = md_path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    try:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass
    style.font.size = Pt(11)

    in_code = False
    code_buf: List[str] = []
    table_buf: List[List[str]] | None = None

    img_re = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")

    def split_table_row(raw: str) -> List[str]:
        s = raw.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]

    i = 0
    while i < len(lines):
        line = lines[i]
        if in_code:
            if line.strip().startswith("```"):
                cp = doc.add_paragraph()
                r = cp.add_run("\n".join(code_buf))
                r.font.name = "Courier New"
                r.font.size = Pt(9)
                cp.paragraph_format.left_indent = Inches(0.2)
                in_code = False
                code_buf = []
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            code_buf = []
            i += 1
            continue

        if line.strip().startswith("|"):
            if table_buf is None:
                table_buf = []
            table_buf.append(split_table_row(line))
            i += 1
            continue
        else:
            if table_buf:
                flush_table(doc, table_buf)
                table_buf = None

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        mimg = img_re.match(stripped)
        if mimg:
            rel = mimg.group(2).strip()
            ip = (base_dir / rel).resolve()
            cap = mimg.group(1).strip() or "图"
            try:
                pc = doc.add_paragraph()
                pc.add_run(cap).bold = True
            except Exception:
                doc.add_paragraph(cap)
            if ip.is_file():
                try:
                    doc.add_picture(str(ip), width=Inches(5.8))
                except Exception as ex:
                    doc.add_paragraph(f"[图片加载失败: {ip} — {ex}]")
            else:
                doc.add_paragraph(f"[缺失图片路径: {ip}]")
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(stripped[2:]).italic = True
        elif stripped.startswith("- "):
            add_paragraph_with_inline(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            add_paragraph_with_inline(
                doc,
                re.sub(r"^\d+\.\s+", "", stripped),
                style="List Number",
            )
        else:
            add_paragraph_with_inline(doc, line.rstrip())

        i += 1

    if table_buf:
        flush_table(doc, table_buf)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print("Saved:", out_path)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("md", type=Path, help="输入 .md")
    ap.add_argument("-o", "--output", type=Path, default=None, help="输出 .docx（默认同名）")
    ap.add_argument(
        "--base-dir",
        type=Path,
        default=None,
        help="解析图片相对路径的根目录（默认同 .md 所在目录）",
    )
    args = ap.parse_args()
    md_path = args.md.expanduser().resolve()
    base = (args.base_dir or md_path.parent).expanduser().resolve()
    out = args.output.expanduser().resolve() if args.output else md_path.with_suffix(".docx")
    convert(md_path, out, base)
